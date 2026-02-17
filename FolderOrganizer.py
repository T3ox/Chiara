import os
import json
import shutil
import base64
import subprocess
import tempfile
import time
from pathlib import Path
import re
from dataclasses import dataclass
from typing import Optional, List, Tuple

from dotenv import load_dotenv
from openai import OpenAI

# Optional libs (installate via pip):
# pip install python-docx openpyxl pillow
from docx import Document
from openpyxl import load_workbook
from PIL import Image


# =========================
# PATHS (la tua struttura)
# =========================
BASE_DIR = Path(__file__).resolve().parent  # cartella dove sta lo script
INPUT_DIR = BASE_DIR / "INPUT_FILES"
RESULT_DIR = BASE_DIR / "RESULT"
CODES_DIR = BASE_DIR  # prompt json stanno qui
ENV_PATH = BASE_DIR / ".env"

# =========================
# PROMPT FILES (nomi esatti dei tuoi json)
# =========================
SUPPORTED_EXT = {
    ".pdf": "prompt_pdf.json",
    ".docx": "prompt_docx.json",
    ".xlsx": "prompt_excel.json",
    ".png": "prompt_immagini.json",
    ".jpg": "prompt_immagini.json",
    ".jpeg": "prompt_immagini.json",
    ".webp": "prompt_immagini.json",
    ".numbers": "prompt_numbers.json",
    ".xls": "prompt_excel.json",
    ".doc": "prompt_docx.json",
}
DEFAULT_PROMPT = "prompt_default.json"

# =========================
# LIMITI (per non mandare tomi)
# =========================
MAX_TEXT_CHARS = 12000          # testo massimo inviato all'API per docx/xlsx convertiti
XLSX_MAX_SHEETS = 4
XLSX_MAX_ROWS_PER_SHEET = 30
XLSX_MAX_COLS_PER_ROW = 12
IMAGE_MAX_SIDE = 1400           # riduzione immagini per invio più leggero


# =========================
# API KEY
# =========================
load_dotenv(dotenv_path=ENV_PATH)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("❌ OPENAI_API_KEY non trovata. Mettila nel file .env nella stessa cartella dello script.")

client = OpenAI(api_key=OPENAI_API_KEY)


# =========================
# UTILS
# =========================
def sanitize_filename(name: str) -> str:
    bad_chars = r'\/:*?"<>|'
    for ch in bad_chars:
        name = name.replace(ch, "")
    # normalizza spazi
    name = name.strip().replace("  ", " ")
    if not name:
        name = "File_Senza_Nome"
    return name

def log_api_call(file_path: Path, mode: str):
    print(f"🌐 API CALL → {file_path.name}  [{mode}]")

def is_file_safe_to_read(file_path: Path) -> tuple[bool, str]:
    if not file_path.exists():
        return False, "file non esiste"
    if not file_path.is_file():
        return False, "non è un file"
    try:
        size = file_path.stat().st_size
    except Exception:
        return False, "impossibile leggere la dimensione"
    if size == 0:
        return False, "file vuoto (0 byte)"
    return True, "ok"


def safe_copy_to_temp_if_locked(file_path: Path) -> Path:
    """
    Se il file è bloccato (WinError 32), lavoriamo su una copia temp.
    """
    try:
        # prova ad aprirlo in read
        with open(file_path, "rb") as f:
            f.read(1)
        return file_path
    except OSError:
        tmp_dir = Path(tempfile.mkdtemp(prefix="folderorganizer_"))
        tmp_path = tmp_dir / file_path.name
        shutil.copy2(file_path, tmp_path)
        return tmp_path
def move_with_retry(src: Path, dst: Path, attempts: int = 6, sleep_s: float = 0.35) -> bool:
    for i in range(attempts):
        try:
            shutil.move(str(src), str(dst))
            return True
        except OSError:
            if i < attempts - 1:
                time.sleep(sleep_s)
            else:
                return False
    return False


# =========================
# RESULT FOLDERS
# =========================
def ensure_result_folders() -> list[str]:
    if not RESULT_DIR.exists():
        raise FileNotFoundError("❌ Cartella RESULT non trovata. Creala e mettici dentro le cartelle di destinazione.")

    # crea DaRevisionare se manca
    da_rev = RESULT_DIR / "DaRevisionare"
    da_rev.mkdir(parents=True, exist_ok=True)

    folders = [f.name for f in RESULT_DIR.iterdir() if f.is_dir()]
    if not folders:
        raise ValueError("❌ Nessuna sottocartella trovata dentro RESULT (anche se DaRevisionare dovrebbe esserci).")

    print("\n📂 Cartelle disponibili in RESULT:")
    for f in folders:
        print(f" - {f}")
    print()

    return folders


# =========================
# PROMPT LOADER + CHECK
# =========================
def check_prompts_exist():
    needed = set(SUPPORTED_EXT.values()) | {DEFAULT_PROMPT}
    missing = [p for p in needed if not (CODES_DIR / p).exists()]
    if missing:
        raise FileNotFoundError(f"❌ Prompt mancanti nella cartella dello script: {missing}")


def load_prompt_for_ext(ext: str, folders: list[str]) -> dict:
    prompt_file = SUPPORTED_EXT.get(ext.lower(), DEFAULT_PROMPT)
    prompt_path = CODES_DIR / prompt_file
    with open(prompt_path, "r", encoding="utf-8") as f:
        prompt = json.load(f)

    # iniezione dinamica cartelle
    prompt.setdefault("user", {})
    prompt["user"]["cartelle_disponibili"] = folders
    return prompt


# =========================
# LOCAL CONVERSIONS
# =========================
def docx_to_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tabelle (se presenti)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells[:XLSX_MAX_COLS_PER_ROW]:
                ct = (cell.text or "").strip().replace("\n", " ")
                if ct:
                    cells.append(ct)
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return text[:MAX_TEXT_CHARS]


def xlsx_to_text(xlsx_path: Path) -> str:
    wb = load_workbook(filename=xlsx_path, data_only=True, read_only=True)
    try:
        sheet_names = wb.sheetnames[:XLSX_MAX_SHEETS]

        out = []
        out.append(f"Workbook: {xlsx_path.name}")
        out.append(f"Sheets: {', '.join(sheet_names)}")
        out.append("")

        for sname in sheet_names:
            ws = wb[sname]
            out.append(f"[SHEET] {sname}")
            for row in ws.iter_rows(min_row=1, max_row=XLSX_MAX_ROWS_PER_SHEET, values_only=True):
                cells = []
                for v in (row[:XLSX_MAX_COLS_PER_ROW] if row else []):
                    cells.append("" if v is None else str(v).strip())
                if any(cells):
                    out.append(" | ".join(cells))
            out.append("")

        text = "\n".join(out)
        return text[:MAX_TEXT_CHARS]

    finally:
        wb.close()

def image_to_data_url(image_path: Path) -> str:
    # riduci dimensione per invio
    img = Image.open(image_path)
    img = img.convert("RGB")
    w, h = img.size
    scale = min(1.0, IMAGE_MAX_SIDE / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)))

    tmp_buf = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
    tmp_buf.close()
    tmp_path = Path(tmp_buf.name)
    img.save(tmp_path, format="JPEG", quality=85, optimize=True)

    b = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)

    b64 = base64.b64encode(b).decode("ascii")
    return f"data:image/jpeg;base64,{b64}"


def can_use_soffice() -> bool:
    # verifica se "soffice" è nel PATH
    try:
        r = subprocess.run(["soffice", "--version"], capture_output=True, text=True, timeout=5)
        return r.returncode == 0
    except Exception:
        return False


def convert_with_soffice(input_path: Path, target_ext: str) -> Path | None:
    """
    Converte con LibreOffice (soffice) se disponibile.
    target_ext es: "docx", "xlsx", "pdf", "txt"
    """
    if not can_use_soffice():
        return None

    out_dir = Path(tempfile.mkdtemp(prefix="soffice_out_"))
    try:
        # --convert-to docx / xlsx / txt / pdf
        r = subprocess.run(
            ["soffice", "--headless", "--convert-to", target_ext, "--outdir", str(out_dir), str(input_path)],
            capture_output=True, text=True, timeout=60
        )
        if r.returncode != 0:
            return None

        # cerca il file convertito
        stem = input_path.stem
        candidates = list(out_dir.glob(f"{stem}*.{target_ext}"))
        if not candidates:
            # in certi casi LibreOffice cambia leggermente il nome
            candidates = list(out_dir.glob(f"*.{target_ext}"))
        return candidates[0] if candidates else None
    except Exception:
        return None


# =========================
# OPENAI CALLS
# =========================
def call_openai_with_pdf(prompt: dict, pdf_path: Path) -> str:
    with open(pdf_path, "rb") as f:
        uploaded = client.files.create(file=f, purpose="assistants")

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=[
            {"role": "system", "content": prompt.get("system", "")},
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": json.dumps(prompt.get("user", {}), ensure_ascii=False)},
                    {"type": "input_file", "file_id": uploaded.id},
                ],
            },
        ],
        max_output_tokens=80
    )
    return response.output_text.strip()

def call_openai_with_text(prompt: dict, extracted_text: str, original_filename: str) -> str:
    # includo anche il nome originale (utile per classificare)
    payload = prompt.get("user", {}).copy()
    payload["nome_file_originale"] = original_filename
    payload["testo_estratto"] = extracted_text

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=[
            {"role": "system", "content": prompt.get("system", "")},
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": json.dumps(payload, ensure_ascii=False)}
                ],
            },
        ],
        max_output_tokens=80
    )
    return response.output_text.strip()


def call_openai_with_image(prompt: dict, image_data_url: str, original_filename: str) -> str:
    payload = prompt.get("user", {}).copy()
    payload["nome_file_originale"] = original_filename

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=[
            {"role": "system", "content": prompt.get("system", "")},
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": json.dumps(payload, ensure_ascii=False)},
                    {"type": "input_image", "image_url": image_data_url},
                ],
            },
        ],
        max_output_tokens=80
    )
    return response.output_text.strip()


def repair_output_format(raw: str, folders: list[str]) -> str:
    """
    Seconda chance ultra rigida: riformatta SOLO nel formato richiesto.
    """
    instruction = {
        "regole": [
            "Riformatta il testo ricevuto nel formato: NuovoNomeDelFile___CartellaDiDestinazione",
            "Usa ESATTAMENTE UNA cartella tra quelle disponibili.",
            "Se la cartella proposta non è nella lista, usa 'DaRevisionare' (se presente) altrimenti la prima cartella.",
            "Nessun altro testo, una sola riga."
        ],
        "cartelle_disponibili": folders,
        "testo_ricevuto": raw
    }

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=[{"role": "user", "content": [{"type": "input_text", "text": json.dumps(instruction, ensure_ascii=False)}]}],
        max_output_tokens=40
    )
    return response.output_text.strip()


# =========================
# OUTPUT VALIDATION
# =========================
def parse_and_validate_result(result: str, folders: list[str]) -> tuple[str, str] | None:
    if not result:
        return None
    if "\n" in result or "\r" in result:
        return None
    if "___" not in result:
        return None

    new_name, folder = result.split("___", 1)
    new_name = new_name.strip()
    folder = folder.strip()

    if not new_name or not folder:
        return None
    if folder not in folders:
        return None
    return new_name, folder


# =========================
# MOVE HELPERS
# =========================
def move_to_result(file_path: Path, folder: str, new_stem: str):
    dest_dir = RESULT_DIR / folder
    dest_dir.mkdir(parents=True, exist_ok=True)

    suffix = file_path.suffix if file_path.suffix else ""
    new_name = sanitize_filename(new_stem) + suffix
    dest_path = dest_dir / new_name

    # duplicati: _1 _2 ...
    if dest_path.exists():
        stem = dest_path.stem
        suf = dest_path.suffix
        i = 1
        while True:
            cand = dest_dir / f"{stem}_{i}{suf}"
            if not cand.exists():
                dest_path = cand
                break
            i += 1

    ok = move_with_retry(file_path, dest_path)
    if not ok:
        print(f"⏭️  Non spostato (file in uso): {file_path.name} — chiudi eventuali anteprime/Explorer e rilancia")
        return
    print(f"✅ Spostato in: {dest_path}")


def move_to_revision(file_path: Path, reason: str = ""):
    folder = "DaRevisionare"
    stem = file_path.stem
    new_stem = f"FAILED_{stem}"
    if reason:
        print(f"⚠️  DaRevisionare: {file_path.name} ({reason})")
    move_to_result(file_path, folder, new_stem)


# =========================
# MAIN PROCESS
# =========================
@dataclass
class RuleHit:
    folder: str
    score: int
    reason: str


def _norm(s: str) -> str:
    s = (s or "").lower()
    # normalizzazione basica
    s = s.replace("\u00a0", " ")
    return s


def _contains_any(hay: str, needles: List[str]) -> bool:
    return any(n in hay for n in needles)


def _regex_any(hay: str, patterns: List[str]) -> bool:
    for p in patterns:
        if re.search(p, hay, flags=re.IGNORECASE | re.MULTILINE):
            return True
    return False


def local_classify(filename: str, ext: str, text: str, available_folders: List[str]) -> Optional[RuleHit]:
    """
    Ritorna RuleHit se trova una classificazione locale affidabile,
    altrimenti None.
    """
    fn = _norm(filename)
    tx = _norm(text)

    hits: List[RuleHit] = []

    # Helper: controlla che una cartella esista davvero nella tua RESULT
    def folder_exists(folder: str) -> bool:
        return folder in available_folders

    # -------------------------
    # A) AMMINISTRAZIONE & FINANZA
    # -------------------------

    # F24
    if _contains_any(fn, ["f24"]) or _regex_any(tx, [r"\bf24\b", r"\bcodice tributo\b", r"\berario\b", r"\binps\b"]):
        folder = "01_AMMINISTRAZIONE_FINANZA/04_F24_Tasse"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 95, "Match F24"))

    # Fatture (attive/passive) - euristica: se c'è "cliente" o nomi noti -> attiva, se "fornitore" -> passiva
    is_invoice = _contains_any(fn, ["fattura", "invoice"]) or _regex_any(tx, [r"\bfattura\b", r"\bimponibile\b", r"\biva\b", r"\btotale\b"])
    if is_invoice:
        # Fornitore / DDT / acquisto -> passiva
        if _contains_any(tx, ["fornitore", "ddt", "bolla", "ordine"]) or _contains_any(fn, ["fornitore", "ddt", "bolla", "ordine"]):
            folder = "01_AMMINISTRAZIONE_FINANZA/02_Fatture_Passive"
            if folder_exists(folder):
                hits.append(RuleHit(folder, 90, "Match fattura passiva (fornitore/ddt/ordine)"))
        else:
            folder = "01_AMMINISTRAZIONE_FINANZA/01_Fatture_Attive"
            if folder_exists(folder):
                hits.append(RuleHit(folder, 85, "Match fattura attiva (pattern imponibile/IVA/totale)"))

    # Liquidazioni IVA / dichiarazioni
    if _contains_any(fn, ["iva", "liquidazione"]) or _regex_any(tx, [r"\bliquidazione iva\b", r"\bcredito iva\b", r"\bdebito iva\b"]):
        folder = "01_AMMINISTRAZIONE_FINANZA/05_IVA_Liquidazioni"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 88, "Match liquidazione IVA"))

    # Estratti conto / banca
    if _contains_any(fn, ["estratto", "conto", "iban", "banca"]) or _regex_any(tx, [r"\bestratto conto\b", r"\bsaldo\b", r"\biban\b"]):
        folder = "01_AMMINISTRAZIONE_FINANZA/03_Banche_Conti"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 82, "Match banca/estratto conto"))

    # Bilanci
    if _contains_any(fn, ["bilancio", "conto economico", "stato patrimoniale"]) or _regex_any(tx, [r"\bconto economico\b", r"\bstato patrimoniale\b"]):
        folder = "01_AMMINISTRAZIONE_FINANZA/06_Bilanci"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 85, "Match bilancio"))

    # -------------------------
    # B) LEGALE / SOCIETARIO / COMPLIANCE
    # -------------------------

    # Visure / registro imprese / CCIAA
    if _contains_any(fn, ["visura", "registro imprese", "camerale", "cciaa"]) or _regex_any(tx, [r"\bregistro imprese\b", r"\bcamera di commercio\b", r"\bvisura\b"]):
        folder = "02_LEGALE_SOCIETARIO/01_Visure_Camerali"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 95, "Match visura/registro imprese"))

    # Contratti
    if _contains_any(fn, ["contratto", "accordo", "scrittura privata"]) or _regex_any(tx, [r"\bcontratto\b", r"\btra le parti\b", r"\boggetto\b", r"\bdecorrenza\b"]):
        # se sembra accordo commerciale / partner
        if _contains_any(tx, ["provvigione", "agente", "collaborazione commerciale", "partner"]) or _contains_any(fn, ["collaborazione", "provvigione"]):
            folder = "02_LEGALE_SOCIETARIO/03_Contratti_Aziendali"
            if folder_exists(folder):
                hits.append(RuleHit(folder, 85, "Match accordo/contratto aziendale"))
        else:
            folder = "02_LEGALE_SOCIETARIO/04_Contratti_Clienti"
            if folder_exists(folder):
                hits.append(RuleHit(folder, 78, "Match contratto (generico, probabile cliente)"))

    # GDPR / privacy
    if _contains_any(fn, ["gdpr", "privacy", "trattamento dati"]) or _regex_any(tx, [r"\bgdpr\b", r"\btrattamento dei dati\b", r"\bresponsabile del trattamento\b"]):
        folder = "02_LEGALE_SOCIETARIO/07_Privacy_GDPR"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 95, "Match GDPR/privacy"))

    # UNI / certificazioni / CEPAS
    if _contains_any(fn, ["uni", "11714", "cepas", "certificazione", "certiquality"]) or _regex_any(tx, [r"\buni\b", r"\b11714\b", r"\bcepas\b", r"\bcertificazione\b"]):
        folder = "02_LEGALE_SOCIETARIO/08_Normative_UNI_Certificazioni"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 92, "Match UNI/certificazioni"))

    # Assicurazioni
    if _contains_any(fn, ["polizza", "assicurazione", "rc"]) or _regex_any(tx, [r"\bpolizza\b", r"\bassicurazione\b", r"\brc\b"]):
        folder = "02_LEGALE_SOCIETARIO/06_Assicurazioni"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 85, "Match assicurazione/polizza"))

    # -------------------------
    # C) CLIENTI / COMMERCIALE
    # -------------------------

    # Preventivi / offerte
    if _contains_any(fn, ["preventivo", "offerta", "proposta"]) or _regex_any(tx, [r"\bpreventivo\b", r"\bproposta\b", r"\bofferta\b"]):
        folder = "03_CLIENTI_COMMERCIALE/02_Preventivi"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 90, "Match preventivo/offerta"))

    # Sopralluoghi
    if _contains_any(fn, ["sopralluogo", "rilievo"]) or _regex_any(tx, [r"\bsopralluogo\b", r"\brilievo\b", r"\bmisure\b"]):
        folder = "03_CLIENTI_COMMERCIALE/05_Sopralluoghi"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 85, "Match sopralluogo/rilievo"))

    # -------------------------
    # D) CANTIERI / OPERATIVO
    # -------------------------

    # SAL
    if _contains_any(fn, ["sal", "stato avanzamento"]) or _regex_any(tx, [r"\bsal\b", r"\bstato avanzamento lavori\b"]):
        folder = "04_CANTIERI_OPERATIVO/06_Stati_Avanzamento_Lavori"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 90, "Match SAL"))

    # Computi metrici
    if _contains_any(fn, ["computo", "mq", "metri quadrati"]) or _regex_any(tx, [r"\bcomputo\b", r"\bquantit", r"\bmq\b"]):
        folder = "04_CANTIERI_OPERATIVO/03_Computi_Metrici"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 82, "Match computo metrico"))

    # Sicurezza cantiere (POS/PSC)
    if _contains_any(fn, ["pos", "psc", "sicurezza"]) or _regex_any(tx, [r"\bpos\b", r"\bpsc\b", r"\bsicurezza\b", r"\brischi\b"]):
        folder = "04_CANTIERI_OPERATIVO/07_Sicurezza_Cantiere"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 92, "Match sicurezza cantiere"))

    # Foto cantieri (solo per immagini)
    if ext in [".png", ".jpg", ".jpeg", ".webp"] and _contains_any(fn, ["cantiere", "prima", "dopo", "lavori"]):
        folder = "04_CANTIERI_OPERATIVO/05_Foto_Cantieri"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 75, "Match foto cantiere da filename"))

    # -------------------------
    # E) HR
    # -------------------------
    if _contains_any(fn, ["busta paga", "cedolino", "stipendio"]) or _regex_any(tx, [r"\bcedolino\b", r"\bbusta paga\b", r"\bnetto\b"]):
        folder = "05_RISORSE_UMANE/04_Buste_Paga"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 95, "Match busta paga/cedolino"))

    if _contains_any(fn, ["contratto lavoro", "assunzione"]) or _regex_any(tx, [r"\bassunzione\b", r"\bcontratto di lavoro\b", r"\blivello\b"]):
        folder = "05_RISORSE_UMANE/03_Contratti_Lavoro"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 90, "Match contratto lavoro/assunzione"))

    # -------------------------
    # F) MARKETING
    # -------------------------
    if _contains_any(fn, ["instagram", "facebook", "post", "locandina", "brochure", "flyer"]) or _contains_any(tx, ["instagram", "facebook"]):
        folder = "06_MARKETING_COMUNICAZIONE/04_Social"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 78, "Match social/marketing"))

    # -------------------------
    # G) PROGETTI SPECIALI
    # -------------------------
    if _contains_any(fn, ["mr white", "mister white", "9010"]) or _contains_any(tx, ["mr white", "mister white", "ral 9010"]):
        folder = "07_PROGETTI_SPECIALI/01_Mr_White"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 92, "Match Mr White"))

    if _contains_any(fn, ["coloraid"]) or _contains_any(tx, ["coloraid"]):
        folder = "07_PROGETTI_SPECIALI/02_ColorAid"
        if folder_exists(folder):
            hits.append(RuleHit(folder, 92, "Match ColorAid"))

    # -------------------------
    # Decide best hit
    # -------------------------
    if not hits:
        return None

    hits.sort(key=lambda x: x.score, reverse=True)
    best = hits[0]

    # Se ci sono due hit ravvicinate (ambiguità), abbassa confidenza
    if len(hits) > 1 and (hits[0].score - hits[1].score) <= 5:
        best = RuleHit(best.folder, max(60, best.score - 10), best.reason + " (ambiguous)")
    return best

def process_one(file_path: Path, folders: list[str]):
    ok, reason = is_file_safe_to_read(file_path)
    if is_windows_file_locked(file_path):
        print(f"⏭️  Skip (file aperto/in uso): {file_path.name} — chiudilo e rilancia")
        return
    if not ok:
        print(f"⏭️  Skip: {file_path.name} ({reason})")
        return

    if not file_path.suffix:
        move_to_revision(file_path, "nessuna estensione")
        return

    ext = file_path.suffix.lower()
    prompt = load_prompt_for_ext(ext, folders)

    print(f"📄 Analizzo: {file_path.name}")
    raw = ""

    try:
        if ext == ".pdf":
            safe_path = safe_copy_to_temp_if_locked(file_path)
            log_api_call(file_path, "PDF")
            raw = call_openai_with_pdf(prompt, safe_path)

        elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
            safe_path = safe_copy_to_temp_if_locked(file_path)
            data_url = image_to_data_url(safe_path)
            log_api_call(file_path, "IMAGE")
            raw = call_openai_with_image(prompt, data_url, file_path.name)

        elif ext == ".docx":
            safe_path = safe_copy_to_temp_if_locked(file_path)
            text = docx_to_text(safe_path)
            if not text.strip():
                move_to_revision(file_path, "docx senza contenuto leggibile")
                return

            hit = local_classify(file_path.name, ext, text, folders)
            if hit and hit.score >= 80:
                move_to_result(file_path, hit.folder, sanitize_filename(file_path.stem))
                print(f"⚡ Classificato LOCAL: {hit.folder} ({hit.score}) - {hit.reason}")
                return

            log_api_call(file_path, "DOCX")
            raw = call_openai_with_text(prompt, text, file_path.name)

        elif ext == ".xlsx":
            safe_path = safe_copy_to_temp_if_locked(file_path)
            text = xlsx_to_text(safe_path)
            if not text.strip():
                move_to_revision(file_path, "xlsx senza contenuto leggibile")
                return

            hit = local_classify(file_path.name, ext, text, folders)
            if hit and hit.score >= 80:
                move_to_result(file_path, hit.folder, sanitize_filename(file_path.stem))
                print(f"⚡ Classificato LOCAL: {hit.folder} ({hit.score}) - {hit.reason}")
                return

            log_api_call(file_path, "XLSX")
            raw = call_openai_with_text(prompt, text, file_path.name)

        elif ext in [".doc", ".xls", ".numbers"]:
            converted = None

            if ext == ".doc":
                converted = convert_with_soffice(file_path, "docx")
                if converted and converted.suffix.lower() == ".docx":
                    text = docx_to_text(converted)
                    if not text.strip():
                        move_to_revision(file_path, "doc->docx ma testo vuoto")
                        return
                    log_api_call(file_path, "DOC→DOCX")
                    raw = call_openai_with_text(prompt, text, file_path.name)
                else:
                    move_to_revision(file_path, "DOC non supportato")
                    return

            elif ext == ".xls":
                converted = convert_with_soffice(file_path, "xlsx")
                if converted and converted.suffix.lower() == ".xlsx":
                    text = xlsx_to_text(converted)
                    if not text.strip():
                        move_to_revision(file_path, "xls->xlsx ma contenuto vuoto")
                        return
                    log_api_call(file_path, "XLS→XLSX")
                    raw = call_openai_with_text(prompt, text, file_path.name)
                else:
                    move_to_revision(file_path, "XLS non supportato")
                    return

            elif ext == ".numbers":
                converted = convert_with_soffice(file_path, "xlsx")
                if converted and converted.suffix.lower() == ".xlsx":
                    text = xlsx_to_text(converted)
                    if not text.strip():
                        move_to_revision(file_path, "numbers->xlsx ma contenuto vuoto")
                        return
                    log_api_call(file_path, "NUMBERS→XLSX")
                    raw = call_openai_with_text(prompt, text, file_path.name)
                else:
                    move_to_revision(file_path, "NUMBERS non supportato")
                    return

        else:
            move_to_revision(file_path, f"estensione non gestita: {ext}")
            return

    except Exception as e:
        move_to_revision(file_path, f"errore lettura/conversione: {e}")
        return

    parsed = parse_and_validate_result(raw, folders)
    if not parsed:
        try:
            fixed = repair_output_format(raw, folders)
            parsed = parse_and_validate_result(fixed, folders)
            if not parsed:
                move_to_revision(file_path, f"output non valido: {raw}")
                return
        except Exception as e:
            move_to_revision(file_path, f"repair fallito: {e}")
            return

    new_stem, folder = parsed
    move_to_result(file_path, folder, new_stem)

def is_windows_file_locked(file_path: Path) -> bool:
    try:
        with open(file_path, "rb") as f:
            f.read(1)
        return False
    except OSError:
        return True

    # Crea RESULT se manca
    RESULT_DIR.mkdir(parents=True, exist_ok=True)

    created = 0
    for rel in folder_tree:
        p = RESULT_DIR / Path(rel)  # Path gestisce "/" anche su Windows
        if not p.exists():
            p.mkdir(parents=True, exist_ok=True)
            created += 1

    print(f"✅ Bootstrap completato. Cartelle create/verificate in RESULT. Nuove create: {created}")


def main():
    if not INPUT_DIR.exists():
        raise FileNotFoundError("❌ Cartella INPUT_FILES non trovata.")

    check_prompts_exist()
    folders = ensure_result_folders()

    files = [p for p in INPUT_DIR.iterdir() if p.is_file()]
    if not files:
        print("📂 Nessun file da processare in INPUT_FILES.")
        return

    for fp in files:
        try:
            process_one(fp, folders)
        except Exception as e:
            # Non deve mai crashare tutto
            print(f"🔥 Errore su {fp.name}: {e}")
            try:
                move_to_revision(fp, f"crash generico: {e}")
            except Exception:
                pass


if __name__ == "__main__":
    main()
